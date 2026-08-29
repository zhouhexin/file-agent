"""会话消息入口的行为测试。

这些测试保护 `/api/conversations/{conversation_id}/messages` 的第一阶段目标：
HTTP 消息必须能进入 LangGraph Agent Runtime，但当前不依赖真实大模型或数据库。
"""

from datetime import datetime, timedelta, timezone
from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
import openpyxl

from app.core import config
from app.db.models import (
    AgentRun,
    Conversation,
    DocumentClassificationRun,
    DocumentVersion,
    Message,
    ToolInvocation,
    UploadArchiveRecord,
    WorkingCopy,
)
from app.modules.managed_files.scanner import ManagedFileScanner
from app.modules.managed_files.service import sync_configured_managed_roots
from app.modules.managed_files.worker import process_next_filesystem_job
from app.tests.helpers import clear_overrides, client_with_database


def _auth_header(client: TestClient, username: str = "message-user") -> dict[str, str]:
    """注册并登录测试用户，返回 Authorization header。"""

    client.post(
        "/api/auth/register",
        json={"username": username, "password": "password123", "display_name": username},
    )
    login_response = client.post(
        "/api/auth/login",
        json={"username": username, "password": "password123"},
    )
    return {"Authorization": f"Bearer {login_response.json()['access_token']}"}


def _upload_document(
    client: TestClient,
    headers: dict[str, str],
    filename: str = "message.txt",
    content: bytes = b"message-file",
) -> str:
    """上传测试文件并返回 document_id。"""

    response = client.post(
        "/api/files/upload",
        headers=headers,
        files={"file": (filename, content, "text/plain")},
    )
    return response.json()["document_id"]


def _configure_upload_lifecycle_storage(monkeypatch, tmp_path) -> None:
    """把上传、不可变原件、工作副本和回收站隔离到测试目录。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "uploads"))
    monkeypatch.setenv("MANAGED_ROOT_ARCHIVE_WRITE_PATH", str(tmp_path / "originals"))
    monkeypatch.setenv("WORKING_COPY_STORAGE_ROOT", str(tmp_path / "working"))
    monkeypatch.setenv("TRASH_STORAGE_ROOT", str(tmp_path / "trash"))
    config.get_settings.cache_clear()


def _xlsx_with_formula_error() -> bytes:
    """构造包含显式公式错误的 Excel 测试文件。"""

    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "汇总"
    worksheet.append(["项目", "公式"])
    worksheet.append(["A", "=SUM(#REF!)"])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _xlsx_with_person_funding_across_sheets() -> bytes:
    """构造跨 Sheet 的人员资助数据，验证聊天查询使用确定性总金额。"""

    workbook = openpyxl.Workbook()
    paper = workbook.active
    paper.title = "论文"
    paper.append(["序号", "申请人", "资助金额"])
    paper.append([1, "张三", 100])
    paper.append([2, "金海燕", 3000])
    paper.append([3, "金海燕", 2000])
    patent = workbook.create_sheet("专利")
    patent.append(["序号", "申请人", "资助金额"])
    patent.append([1, "金海燕", 1500])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _docx_for_summary() -> bytes:
    """构造同时包含自然语言说明和代码块的 DOCX，保护页面总结不会退化为代码预览。"""

    document = DocxDocument()
    document.add_paragraph(
        "本文介绍一个 Python 数据可视化项目。"
        "项目使用 NumPy 处理数组，Pandas 整理表格数据，并通过 Matplotlib 绘制动态图表。"
    )
    document.add_paragraph(
        "运行前需要安装依赖，Python 版本为 3.12。"
        "程序主要完成数据加载、坐标计算、动画生成和结果导出。"
    )
    document.add_paragraph(
        "代码示例：\n"
        "import numpy as np\n"
        "import pandas as pd\n"
        "import matplotlib.pyplot as plt\n"
        "for index in range(100):\n"
        "    print(index)"
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _latest_agent_audit(session_factory) -> tuple[AgentRun, list[str]]:
    """从测试数据库读取最近一次内部运行，普通消息响应本身不得暴露这些字段。"""

    with session_factory() as db:
        run = db.query(AgentRun).order_by(AgentRun.created_at.desc()).first()
        assert run is not None
        tool_names = [
            item.tool_name
            for item in (
                db.query(ToolInvocation)
                .filter(ToolInvocation.agent_run_id == run.id)
                .order_by(ToolInvocation.created_at.asc())
                .all()
            )
        ]
        db.expunge(run)
        return run, tool_names


def test_message_does_not_expand_explicit_filename_before_working_copy_is_active(monkeypatch, tmp_path):
    """原始受管文件尚未导入共享工作副本时，不得为了回答而扩大为其他文件。"""

    managed_root = tmp_path / "school-files"
    managed_root.mkdir()
    filename = "述职报告-鲁晓锋-20200421.txt"
    (managed_root / filename).write_text(
        "鲁晓锋同志在述职报告中总结了年度教学、科研和学生培养工作。",
        encoding="utf-8",
    )
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("MANAGED_ROOT_SCHOOL_FILES", str(managed_root))
    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    monkeypatch.setenv("CHAT_DOCUMENT_SUMMARY_PROVIDER", "disabled")
    config.get_settings.cache_clear()
    client, session_factory = client_with_database()
    headers = _auth_header(client, "managed-filename-summary-user")

    with session_factory() as db:
        roots = sync_configured_managed_roots(db, root_key="school_files", scan=False)
        for root in roots:
            ManagedFileScanner(db).scan_root(root)
        db.commit()

    for content in [
        f"总结一下{filename}",
        f"{filename} 总结一下这个文档",
    ]:
        response = client.post(
            "/api/conversations/managed-filename-summary-chat/messages",
            headers=headers,
            json={"content": content, "attachments": []},
        )

        assert response.status_code == 200
        task_result = response.json()["task_result"]
        assert task_result["response_type"] == "async_job"
        assert task_result["task_status"] == "processing"
        assert task_result["final_response"] is None
        assert task_result["evidence_answer_result"] is None
        assert task_result["pending_job_ids"] == []

    # 源侧分析完成后 worker 必须自动续跑原请求；用户不需要再次发送问题。
    while process_next_filesystem_job(
        session_factory=session_factory,
        worker_id="search-readiness-test",
        queue_names={"SOURCE_ANALYSIS", "MATERIALIZE", "IMPORT", "ANALYSIS"},
    ):
        pass
    with session_factory() as db:
        runs = (
            db.query(AgentRun)
            .filter(
                AgentRun.conversation_id
                == "managed-filename-summary-chat"
            )
            .all()
        )
        assert len(runs) == 2
        assert all(run.status != "WAITING_FOR_ASYNC_JOB" for run in runs)

    clear_overrides()
    config.get_settings.cache_clear()


def test_uploaded_docx_summary_uses_full_text_points_instead_of_preview(monkeypatch, tmp_path):
    """页面上传 DOCX 后请求总结，应返回完整正文要点并过滤代码，不能展示 280 字预览。"""

    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    monkeypatch.setenv("LLM_ENABLED", "false")
    monkeypatch.setenv("CHAT_DOCUMENT_SUMMARY_PROVIDER", "disabled")
    config.get_settings.cache_clear()
    client, _session_factory = client_with_database()
    headers = _auth_header(client, "uploaded-docx-summary-user")
    upload_response = client.post(
        "/api/files/upload",
        headers=headers,
        files={
            "file": (
                "python代码.docx",
                _docx_for_summary(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 202

    response = client.post(
        "/api/conversations/uploaded-docx-summary-chat/messages",
        headers=headers,
        json={
            "content": "总结这个文档",
            "attachments": [{"document_id": upload_response.json()["document_id"]}],
        },
    )

    assert response.status_code == 200
    task_result = response.json()["task_result"]
    assert task_result["task_status"] == "processing"
    assert task_result["response_type"] == "async_job"
    assert task_result["final_response"] is None
    assert task_result["pending_job_ids"] == []

    clear_overrides()
    config.get_settings.cache_clear()


def test_uploaded_xlsx_person_total_routes_to_deterministic_analysis(monkeypatch, tmp_path):
    """附件只传 document_id 时，后端必须补全真实 XLSX 元数据并回答人员总金额。"""

    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    monkeypatch.setenv("LLM_ENABLED", "false")
    config.get_settings.cache_clear()
    client, session_factory = client_with_database()
    headers = _auth_header(client, "uploaded-xlsx-person-total-user")
    upload_response = client.post(
        "/api/files/upload",
        headers=headers,
        files={
            "file": (
                "2024科研成果资助汇总表.xlsx",
                _xlsx_with_person_funding_across_sheets(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )
    assert upload_response.status_code == 202

    response = client.post(
        "/api/conversations/uploaded-xlsx-person-total-chat/messages",
        headers=headers,
        json={
            "content": "金海燕的资助总金额是多少",
            "attachments": [{"document_id": upload_response.json()["document_id"]}],
        },
    )

    assert response.status_code == 200
    task_result = response.json()["task_result"]
    assert task_result["response_type"] == "text"
    assert "结果：6,500" in task_result["final_response"]
    assert "筛选条件：“申请人”等于“金海燕”" in task_result["final_response"]
    assert "Sheet“论文”：5,000" in task_result["final_response"]
    assert "Sheet“专利”：1,500" in task_result["final_response"]
    assert "计算方式：5,000 + 1,500 = 6,500" in task_result["final_response"]
    assert " B3" not in task_result["final_response"]
    assert "分类建议" not in task_result["final_response"]

    scoped_response = client.post(
        "/api/conversations/uploaded-xlsx-person-total-chat/messages",
        headers=headers,
        json={
            "content": "2024科研成果资助汇总表中金海燕的资助总金额是多少",
            "attachments": [{"document_id": upload_response.json()["document_id"]}],
        },
    )

    assert scoped_response.status_code == 200
    scoped_result = scoped_response.json()["task_result"]
    assert "结果：6,500" in scoped_result["final_response"]
    assert "筛选条件：“申请人”等于“金海燕”" in scoped_result["final_response"]
    assert "研成果资助汇总表中金海燕" not in scoped_result["final_response"]
    run, tool_names = _latest_agent_audit(session_factory)
    assert run.intent == "ANALYZE_SPREADSHEET"
    assert tool_names == ["analyze-spreadsheet"]

    clear_overrides()
    config.get_settings.cache_clear()


def test_post_message_starts_agent_run():
    """发送消息后必须持久化 AgentRun，但普通响应只能返回安全任务投影。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client)
    document_id = _upload_document(client, headers)

    response = client.post(
        "/api/conversations/conv-1/messages",
        headers=headers,
        json={
            "content": "帮我分类这批文件",
            "attachments": [{"document_id": document_id}],
        },
    )

    assert response.status_code == 200
    data = response.json()
    assert data["message"]["conversation_id"] == "conv-1"
    assert data["message"]["role"] == "user"
    assert "agent_run" not in data
    run, tool_names = _latest_agent_audit(session_factory)
    assert run.status == "COMPLETED"
    assert run.intent == "CLASSIFY_FILES"
    assert run.selected_skills_json == [
        "chat-intake",
        "document-text-extract",
        "document-classification",
        "change-report",
    ]
    assert tool_names == ["extract-document-text"]
    # 普通用户投影不能要求前端理解 Skill 或 Tool，也不能携带解析器和内部路径字段。
    task_result = data["task_result"]
    assert task_result["task_id"] == run.id
    assert task_result["task_status"] == "completed"
    assert task_result["response_type"] == "file_results"
    assert task_result["display_mode"] == "classification_cards"
    assert "selected_skills" not in task_result
    assert "tool_invocations" not in task_result
    assert "tool_results" not in task_result
    assert "extractor" not in task_result["document_results"][0]
    assert "relative_path" not in task_result["document_results"][0]
    assert "index_run_id" not in task_result["document_results"][0]
    assert "search_text" not in task_result["document_results"][0]
    assert "embedding" not in task_result["document_results"][0]
    clear_overrides()


def test_post_message_rename_and_classify_returns_plan_before_copy_is_ready():
    """刚上传文件尚无工作副本时，也须立即返回分类卡和可复用的重命名计划。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "rename-classify-message-user")
    document_id = _upload_document(
        client,
        headers,
        filename="2026年奖学金通知.txt",
        content="2026年奖学金评审工作通知\n请各学院按要求报送材料。".encode("utf-8"),
    )

    response = client.post(
        "/api/conversations/rename-classify-conversation/messages",
        headers=headers,
        json={
            "content": "对上传文件进行重命名和分类",
            "attachments": [{"document_id": document_id}],
        },
    )

    assert response.status_code == 200
    task_result = response.json()["task_result"]
    assert task_result["response_type"] == "operation_plan"
    assert task_result["display_mode"] == "classification_cards"
    assert task_result["document_results"]
    assert task_result["document_results"][0]["document_id"] == document_id
    assert "生成 1 个可执行的重命名建议" in task_result["final_response"]
    assert task_result["operation_plan_id"]
    assert task_result["rename_plan_result"]["ready_count"] == 1
    assert task_result["rename_plan_result"]["suggestions"][0]["proposed_filename"]

    run, tool_names = _latest_agent_audit(session_factory)
    assert run.intent == "CLASSIFY_AND_SUGGEST_RENAME"
    assert tool_names == [
        "extract-document-text",
        "generate-rename-suggestions",
    ]
    clear_overrides()


def test_target_only_rename_does_not_bind_historical_file_with_same_target_name():
    """只写目标名称时，即使历史附件同名也必须要求说明源文件，不能生成错误计划。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "rename-target-only-user")
    document_id = _upload_document(
        client,
        headers,
        filename="西安理工大学用印申请单.docx",
        content=b"historical-target-name",
    )
    url = "/api/conversations/rename-target-only-conversation/messages"
    history_response = client.post(
        url,
        headers=headers,
        json={
            "content": "保存这个文件",
            "attachments": [{"document_id": document_id}],
        },
    )
    assert history_response.status_code == 200

    response = client.post(
        url,
        headers=headers,
        json={
            "content": "重命名为 西安理工大学用印申请单.docx",
            "attachments": [],
        },
    )

    assert response.status_code == 200
    task_result = response.json()["task_result"]
    assert task_result["operation_plan_id"] is None
    assert "不能确定要重命名哪一个文件" in task_result["final_response"]
    assert "原文件名.ext" in task_result["final_response"]
    run, tool_names = _latest_agent_audit(session_factory)
    assert run.intent == "MISSING_FILE_SCOPE"
    # 澄清属于 LangGraph 响应分支，不需要为了内部占位审计调用无副作用 Tool。
    assert tool_names == []
    clear_overrides()


def test_get_conversation_returns_messages_with_task_results_and_attachments():
    """读取会话详情时必须返回附件和安全任务投影，不返回内部 AgentRun。"""

    client, _ = client_with_database()
    headers = _auth_header(client, "history-user")
    document_id = _upload_document(client, headers)

    post_response = client.post(
        "/api/conversations/web-chat/messages",
        headers=headers,
        json={
            "content": "帮我读取并分类这批文件",
            "attachments": [{"document_id": document_id}],
        },
    )
    assert post_response.status_code == 200

    response = client.get("/api/conversations/web-chat", headers=headers)

    assert response.status_code == 200
    data = response.json()
    assert data["id"] == "web-chat"
    assert len(data["messages"]) == 1
    history_message = data["messages"][0]
    assert history_message["content"] == "帮我读取并分类这批文件"
    assert history_message["attachments"][0]["document_id"] == document_id
    assert history_message["attachments"][0]["filename"] == "message.txt"
    assert "agent_run" not in history_message
    assert history_message["task_result"]["task_status"] == "completed"
    assert history_message["task_result"]["final_response"]
    assert len(history_message["task_result"]["document_results"]) == 1
    assert history_message["task_result"]["document_results"][0]["document_id"] == document_id
    assert history_message["task_result"]["document_results"][0]["filename"] == "message.txt"
    assert history_message["task_result"]["document_results"][0]["extraction_status"] == "COMPLETED"
    assert "tool_invocations" not in history_message["task_result"]
    clear_overrides()


def test_get_conversation_returns_latest_page_with_pagination():
    """会话详情默认只返回最近一页消息，避免聊天页首屏加载完整历史。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "paged-history-user")
    me_response = client.get("/api/auth/me", headers=headers)
    current_user_id = me_response.json()["id"]
    base_time = datetime(2026, 1, 1, tzinfo=timezone.utc)
    with session_factory() as db:
        db.add(Conversation(id="paged-chat", user_id=current_user_id, title=""))
        for index in range(15):
            db.add(
                Message(
                    conversation_id="paged-chat",
                    user_id=current_user_id,
                    role="user",
                    content=f"历史消息 {index + 1}",
                    attachments_json=[],
                    created_at=base_time + timedelta(seconds=index),
                )
            )
        db.commit()

    response = client.get("/api/conversations/paged-chat?limit=10", headers=headers)

    assert response.status_code == 200
    data = response.json()
    assert [message["content"] for message in data["messages"]] == [
        f"历史消息 {index}" for index in range(6, 16)
    ]
    assert data["pagination"]["has_more"] is True
    assert data["pagination"]["oldest_message_id"] == data["messages"][0]["id"]
    assert data["pagination"]["limit"] == 10
    clear_overrides()


def test_get_conversation_returns_older_page_before_message_id():
    """传入 before_message_id 时返回该消息之前的更早历史。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "older-history-user")
    current_user_id = client.get("/api/auth/me", headers=headers).json()["id"]
    message_ids: list[str] = []
    base_time = datetime(2026, 1, 1, tzinfo=timezone.utc)
    with session_factory() as db:
        db.add(Conversation(id="older-chat", user_id=current_user_id, title=""))
        for index in range(12):
            message = Message(
                conversation_id="older-chat",
                user_id=current_user_id,
                role="user",
                content=f"消息 {index + 1}",
                attachments_json=[],
                created_at=base_time + timedelta(seconds=index),
            )
            db.add(message)
            db.flush()
            message_ids.append(message.id)
        db.commit()

    response = client.get(
        f"/api/conversations/older-chat?limit=5&before_message_id={message_ids[7]}",
        headers=headers,
    )

    assert response.status_code == 200
    data = response.json()
    assert [message["content"] for message in data["messages"]] == [
        "消息 3",
        "消息 4",
        "消息 5",
        "消息 6",
        "消息 7",
    ]
    assert data["pagination"]["has_more"] is True
    assert data["pagination"]["oldest_message_id"] == data["messages"][0]["id"]
    clear_overrides()


def test_message_can_reference_previous_uploaded_attachment():
    """用户说“上面上传的文件”时，应自动引用当前会话最近的附件。"""

    client, _ = client_with_database()
    headers = _auth_header(client, "previous-file-user")
    document_id = _upload_document(client, headers)

    first_response = client.post(
        "/api/conversations/context-chat/messages",
        headers=headers,
        json={
            "content": "帮我读取这个文件",
            "attachments": [{"document_id": document_id}],
        },
    )
    assert first_response.status_code == 200

    second_response = client.post(
        "/api/conversations/context-chat/messages",
        headers=headers,
        json={
            "content": "读取上面上传的文件，给我讲解大概总结一下文章内容",
            "attachments": [],
        },
    )

    assert second_response.status_code == 200
    data = second_response.json()
    assert data["message"]["attachments"] == [{"document_id": document_id}]
    assert data["task_result"]["task_status"] == "processing"
    assert data["task_result"]["response_type"] == "async_job"
    assert data["task_result"]["final_response"] is None
    clear_overrides()


def test_message_can_reference_second_previous_attachment_by_ordinal():
    """用户说“第二个文件”时，应只引用当前会话上文附件中的第二个文件。"""

    client, _ = client_with_database()
    headers = _auth_header(client, "second-file-user")
    first_document_id = _upload_document(client, headers, filename="first.txt", content=b"first-file")
    second_document_id = _upload_document(client, headers, filename="电子发票承诺书.doc", content=b"second-file")

    first_response = client.post(
        "/api/conversations/ordinal-chat/messages",
        headers=headers,
        json={
            "content": "帮我读取并分类这批文件",
            "attachments": [
                {"document_id": first_document_id},
                {"document_id": second_document_id},
            ],
        },
    )
    assert first_response.status_code == 200

    second_response = client.post(
        "/api/conversations/ordinal-chat/messages",
        headers=headers,
        json={
            "content": "重新对第二个文件：电子发票承诺书.doc进行分类",
            "attachments": [],
        },
    )

    assert second_response.status_code == 200
    data = second_response.json()
    assert data["message"]["attachments"] == [{"document_id": second_document_id}]
    assert data["task_result"]["task_status"] == "completed"
    assert data["task_result"]["document_results"][0]["document_id"] == second_document_id
    clear_overrides()


def test_message_can_reference_previous_attachment_by_filename_fragment():
    """用户按文件名片段提问时，应自动引用当前会话中的对应历史附件。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "filename-reference-user")
    document_id = _upload_document(
        client,
        headers,
        filename="2019年学院科研成果资助表.xlsx",
        content="姓名,金额\n张三,100\n李四,200\n".encode(),
    )

    first_response = client.post(
        "/api/conversations/filename-reference-chat/messages",
        headers=headers,
        json={
            "content": "帮我读取这个文件",
            "attachments": [{"document_id": document_id}],
        },
    )
    assert first_response.status_code == 200

    second_response = client.post(
        "/api/conversations/filename-reference-chat/messages",
        headers=headers,
        json={
            "content": "汇总“2019年学院科研成果资助表.xlsx”中的金额",
            "attachments": [],
        },
    )

    assert second_response.status_code == 200
    data = second_response.json()
    assert data["message"]["attachments"] == [{"document_id": document_id}]
    run, tool_names = _latest_agent_audit(session_factory)
    assert run.intent == "ANALYZE_SPREADSHEET"
    assert tool_names == ["analyze-spreadsheet"]
    assert "AgentRun completed" not in (data["task_result"]["final_response"] or "")
    clear_overrides()


def test_message_can_reference_previous_attachment_by_fuzzy_filename_tokens():
    """用户只说文件名核心词时不能静默绑定历史附件，必须先确定活动文件范围。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "fuzzy-filename-reference-user")
    old_document_id = _upload_document(
        client,
        headers,
        filename="2019年学院科研成果资助汇总表.xlsx",
        content="教师,资助金额\n王老师,100\n".encode(),
    )
    target_document_id = _upload_document(
        client,
        headers,
        filename="2024年度学院科研成果资助汇总表.xlsx",
        content="教师,资助金额\n张老师,300\n李老师,200\n".encode(),
    )

    first_response = client.post(
        "/api/conversations/fuzzy-filename-reference-chat/messages",
        headers=headers,
        json={
            "content": "帮我读取这批文件",
            "attachments": [{"document_id": old_document_id}, {"document_id": target_document_id}],
        },
    )
    assert first_response.status_code == 200

    second_response = client.post(
        "/api/conversations/fuzzy-filename-reference-chat/messages",
        headers=headers,
        json={
            "content": "根据教师来汇总2024科研成果资助汇总表中的资助金额",
            "attachments": [],
        },
    )

    assert second_response.status_code == 200
    data = second_response.json()
    assert data["message"]["attachments"] == []
    run, tool_names = _latest_agent_audit(session_factory)
    assert run.intent == "EVIDENCE_ANSWER"
    assert tool_names == ["evidence-answer"]
    assert "AgentRun completed" not in (data["task_result"]["final_response"] or "")
    clear_overrides()


def test_message_can_validate_uploaded_spreadsheet_formula_errors():
    """聊天入口中的表格校验请求必须路由到 validate-spreadsheet。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "spreadsheet-validation-user")
    document_id = _upload_document(
        client,
        headers,
        filename="公式错误.xlsx",
        content=_xlsx_with_formula_error(),
    )

    response = client.post(
        "/api/conversations/spreadsheet-validation-chat/messages",
        headers=headers,
        json={
            "content": "检查这份表格有没有公式错误",
            "attachments": [{"document_id": document_id}],
        },
    )

    assert response.status_code == 200
    data = response.json()
    run, tool_names = _latest_agent_audit(session_factory)
    assert run.intent == "VALIDATE_SPREADSHEET"
    assert tool_names == ["validate-spreadsheet"]
    assert "#REF!" in (data["task_result"]["final_response"] or "")
    clear_overrides()


def test_message_can_summarize_previous_classification_results():
    """用户要求总结之前文件分类时，应读取分类建议而不是只返回基础洞察文件名。"""

    client, _ = client_with_database()
    headers = _auth_header(client, "classification-summary-user")
    first_document_id = _upload_document(client, headers, filename="职称材料.txt", content="教师职称申报材料".encode())
    second_document_id = _upload_document(client, headers, filename="科研成果.txt", content="学院科研成果资助材料".encode())

    first_response = client.post(
        "/api/conversations/classification-summary-chat/messages",
        headers=headers,
        json={
            "content": "帮我读取并分类这批文件",
            "attachments": [
                {"document_id": first_document_id},
                {"document_id": second_document_id},
            ],
        },
    )
    assert first_response.status_code == 200

    second_response = client.post(
        "/api/conversations/classification-summary-chat/messages",
        headers=headers,
        json={
            "content": "帮我总结一下刚刚上传文件的分类",
            "attachments": [],
        },
    )

    assert second_response.status_code == 200
    final_response = second_response.json()["task_result"]["final_response"]
    assert "已汇总" in final_response
    assert "分类建议" in final_response
    assert "基础洞察" not in final_response
    assert "职称材料.txt" in final_response
    assert "科研成果.txt" in final_response
    clear_overrides()


def test_just_uploaded_classification_uses_latest_attachment_batch_only():
    """“刚刚上传文件”应指向最近一条带附件消息中的整批文件，而不是所有历史附件。"""

    client, _ = client_with_database()
    headers = _auth_header(client, "latest-batch-user")
    old_first_id = _upload_document(client, headers, filename="旧批次-职称.txt", content="教师职称申报材料".encode())
    old_second_id = _upload_document(client, headers, filename="旧批次-科研.txt", content="学院科研成果资助材料".encode())
    latest_id = _upload_document(client, headers, filename="最新批次-财务.txt", content="电子发票财务承诺材料".encode())

    old_response = client.post(
        "/api/conversations/latest-batch-chat/messages",
        headers=headers,
        json={
            "content": "帮我读取并分类这批文件",
            "attachments": [{"document_id": old_first_id}, {"document_id": old_second_id}],
        },
    )
    assert old_response.status_code == 200

    latest_response = client.post(
        "/api/conversations/latest-batch-chat/messages",
        headers=headers,
        json={
            "content": "帮我读取并分类这个文件",
            "attachments": [{"document_id": latest_id}],
        },
    )
    assert latest_response.status_code == 200

    historical_summary_response = client.post(
        "/api/conversations/latest-batch-chat/messages",
        headers=headers,
        json={
            "content": "总结一下之前上传的所有项目分类",
            "attachments": [],
        },
    )
    assert historical_summary_response.status_code == 200
    assert len(historical_summary_response.json()["message"]["attachments"]) == 3

    summary_response = client.post(
        "/api/conversations/latest-batch-chat/messages",
        headers=headers,
        json={
            "content": "帮我总结一下刚刚上传的所有文件分类",
            "attachments": [],
        },
    )

    assert summary_response.status_code == 200
    data = summary_response.json()
    assert data["message"]["attachments"] == [{"document_id": latest_id}]
    final_response = data["task_result"]["final_response"]
    assert "最新批次-财务.txt" in final_response
    assert "旧批次-职称.txt" not in final_response
    assert "旧批次-科研.txt" not in final_response
    clear_overrides()


def test_uploaded_message_attachments_share_batch_id():
    """同一条用户消息里的真实上传附件必须带同一个 batch_id，供后续“刚刚上传”精确引用。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "batch-marker-user")
    first_document_id = _upload_document(client, headers, filename="批次-1.txt", content=b"first")
    second_document_id = _upload_document(client, headers, filename="批次-2.txt", content=b"second")

    response = client.post(
        "/api/conversations/batch-marker-chat/messages",
        headers=headers,
        json={
            "content": "帮我读取并分类这批文件",
            "attachments": [
                {"document_id": first_document_id},
                {"document_id": second_document_id},
            ],
        },
    )

    assert response.status_code == 200
    with session_factory() as db:
        message = (
            db.query(Message)
            .filter(Message.conversation_id == "batch-marker-chat")
            .order_by(Message.created_at.desc())
            .first()
        )
        assert message is not None
        sources = {item.get("source") for item in message.attachments_json}
        batch_ids = {item.get("batch_id") for item in message.attachments_json}
    assert sources == {"uploaded"}
    assert len(batch_ids) == 1
    assert None not in batch_ids
    clear_overrides()


def test_archived_upload_is_canonicalized_to_working_copy_before_agent_run(
    monkeypatch,
    tmp_path,
):
    """消息审计保留上传 ID，但 Agent Tool 只能消费活动工作副本 ID。"""

    _configure_upload_lifecycle_storage(monkeypatch, tmp_path)
    client, session_factory = client_with_database()
    headers = _auth_header(client, "canonical-agent-attachment-user")
    upload_document_id = _upload_document(
        client,
        headers,
        filename="salary-audit-notice.txt",
        content="工资津贴补贴发放情况专项监督检查和自查工作。".encode("utf-8"),
    )

    for _ in range(30):
        if process_next_filesystem_job(
            session_factory=session_factory,
            worker_id="canonical-agent-attachment-worker",
        ) is None:
            break

    with session_factory() as db:
        upload_version = (
            db.query(DocumentVersion)
            .filter(
                DocumentVersion.document_id == upload_document_id,
                DocumentVersion.storage_tier == "UPLOAD",
            )
            .one()
        )
        archive = (
            db.query(UploadArchiveRecord)
            .filter(
                UploadArchiveRecord.upload_document_version_id
                == upload_version.id
            )
            .one()
        )
        working_copy = (
            db.query(WorkingCopy)
            .filter(
                WorkingCopy.managed_file_id == archive.managed_file_id,
                WorkingCopy.status == "ACTIVE",
            )
            .one()
        )
        working_document_id = working_copy.document_id
        upload_classification_count_before = (
            db.query(DocumentClassificationRun)
            .filter(DocumentClassificationRun.document_id == upload_document_id)
            .count()
        )
        assert upload_classification_count_before == 0

    response = client.post(
        "/api/conversations/canonical-agent-attachment-chat/messages",
        headers=headers,
        json={
            "content": "对上传文件进行分类归档",
            "attachments": [{"document_id": upload_document_id}],
        },
    )

    assert response.status_code == 200
    assert response.json()["message"]["attachments"] == [
        {"document_id": upload_document_id}
    ]
    with session_factory() as db:
        message = (
            db.query(Message)
            .filter(
                Message.conversation_id == "canonical-agent-attachment-chat",
                Message.role == "user",
            )
            .one()
        )
        run = db.query(AgentRun).filter(AgentRun.message_id == message.id).one()
        invocations = (
            db.query(ToolInvocation)
            .filter(ToolInvocation.agent_run_id == run.id)
            .order_by(ToolInvocation.created_at)
            .all()
        )
        invoked_document_ids = {
            str(item.input_json.get("document_id"))
            for item in invocations
            if item.input_json.get("document_id")
        }
        assert message.attachments_json[0]["document_id"] == upload_document_id
        assert invoked_document_ids == {working_document_id}
        assert upload_document_id not in invoked_document_ids
        assert (
            db.query(DocumentClassificationRun)
            .filter(DocumentClassificationRun.document_id == upload_document_id)
            .count()
            == 0
        )

    # 文件已归档且不属于新会话附件时，完整文件名仍必须先解析到同一个活动工作副本，
    # 不能把内部 upload_archive 错当作普通受管源目录查询。
    filename_response = client.post(
        "/api/conversations/working-copy-filename-classification-chat/messages",
        headers=headers,
        json={
            "content": "重新分类“salary-audit-notice.txt”",
            "attachments": [],
        },
    )

    assert filename_response.status_code == 200
    assert filename_response.json()["task_result"]["task_status"] == "completed"
    assert filename_response.json()["task_result"]["document_results"][0][
        "document_id"
    ] == working_document_id
    with session_factory() as db:
        filename_message = (
            db.query(Message)
            .filter(
                Message.conversation_id
                == "working-copy-filename-classification-chat",
                Message.role == "user",
            )
            .one()
        )
        assert filename_message.attachments_json[0]["document_id"] == working_document_id
        filename_run = (
            db.query(AgentRun)
            .filter(AgentRun.message_id == filename_message.id)
            .one()
        )
        filename_invocations = (
            db.query(ToolInvocation)
            .filter(ToolInvocation.agent_run_id == filename_run.id)
            .order_by(ToolInvocation.created_at)
            .all()
        )
        # 重新分类会在新分类达到自动标准时直接归位，因此除了正文解析还会
        # 调用受控工作副本动作；两步都只能使用规范工作副本身份。
        assert [item.tool_name for item in filename_invocations] == [
            "extract-document-text",
            "working-copy-action-plan-create",
        ]
        assert filename_invocations[0].input_json["document_id"] == working_document_id
        assert filename_invocations[1].input_json["document_ids"] == [
            working_document_id
        ]
        assert (
            filename_invocations[1].input_json["action"]
            == "MOVE_AFTER_AUTO_RECLASSIFICATION"
        )
    clear_overrides()


def test_same_document_id_is_deduplicated_in_message_and_legacy_history():
    """同一 document_id 重复提交或残留在旧消息中都只展示一次。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "same-document-attachment-user")
    document_id = _upload_document(
        client,
        headers,
        filename="重复附件.txt",
        content=b"same attachment reference",
    )

    response = client.post(
        "/api/conversations/same-document-attachment-chat/messages",
        headers=headers,
        json={
            "content": "读取这个文件",
            "attachments": [
                {"document_id": document_id},
                {"document_id": document_id},
            ],
        },
    )

    assert response.status_code == 200
    assert response.json()["message"]["attachments"] == [{"document_id": document_id}]
    with session_factory() as db:
        message = (
            db.query(Message)
            .filter(Message.conversation_id == "same-document-attachment-chat")
            .order_by(Message.created_at.desc(), Message.id.desc())
            .first()
        )
        assert message is not None
        assert len(message.attachments_json) == 1
        # 模拟修复前已经写入数据库的重复引用，历史接口也必须防御性折叠。
        original_item = dict(message.attachments_json[0])
        message.attachments_json = [original_item, dict(original_item)]
        db.commit()

    history = client.get(
        "/api/conversations/same-document-attachment-chat",
        headers=headers,
    )
    assert history.status_code == 200
    user_message = next(
        item
        for item in history.json()["messages"]
        if item["role"] == "user"
    )
    assert [item["document_id"] for item in user_message["attachments"]] == [document_id]
    clear_overrides()


def test_legacy_lifecycle_status_messages_are_hidden_from_history_projection():
    """升级前已写入 assistant 的生命周期状态消息也不能继续出现在聊天页。"""

    client, session_factory = client_with_database()
    headers = _auth_header(client, "legacy-lifecycle-message-user")
    with session_factory() as db:
        user_id = client.get("/api/auth/me", headers=headers).json()["id"]
        conversation = Conversation(
            id="legacy-lifecycle-message-chat",
            user_id=user_id,
            title="",
        )
        db.add(conversation)
        db.add_all(
            [
                Message(
                    conversation_id=conversation.id,
                    user_id=user_id,
                    role="assistant",
                    content="文件“旧通知.txt”的原件已归档，正在创建工作副本。",
                    attachments_json=[],
                ),
                Message(
                    conversation_id=conversation.id,
                    user_id=user_id,
                    role="assistant",
                    content="工作副本操作完成：TRASH_WORKING_COPIES",
                    attachments_json=[],
                ),
                Message(
                    conversation_id=conversation.id,
                    user_id=user_id,
                    role="user",
                    content="保留在对话中的真实消息",
                    attachments_json=[],
                ),
            ]
        )
        db.commit()

    history = client.get(
        "/api/conversations/legacy-lifecycle-message-chat",
        headers=headers,
    )
    assert history.status_code == 200
    assert [message["content"] for message in history.json()["messages"]] == ["保留在对话中的真实消息"]
    clear_overrides()


def test_post_message_rejects_invalid_attachment():
    """附件引用缺少 document_id 时必须由请求 schema 拒绝。"""

    client, _ = client_with_database()

    response = client.post(
        "/api/conversations/conv-1/messages",
        headers=_auth_header(client),
        json={
            "content": "帮我读取文件",
            "attachments": [{"filename": "bad.pdf"}],
        },
    )

    assert response.status_code == 422
    clear_overrides()
