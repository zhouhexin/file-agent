"""旧版 Office 派生件转换与复用测试。"""

from __future__ import annotations

import hashlib
import os
from pathlib import Path, PureWindowsPath
import subprocess

from docx import Document as DocxDocument
import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.db.base import Base
from app.db.models import Document, DocumentArtifact, DocumentVersion
from app.modules.files.office_conversion import (
    LegacyOfficeConversionService,
    OfficeConversionError,
    _replace_with_retry,
    libreoffice_profile_uri,
    resolve_libreoffice_executable,
)


def _session():
    """创建包含完整 ORM 表的隔离数据库会话。"""

    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    return sessionmaker(bind=engine, autoflush=False, autocommit=False)()


def _document(*, source_path: Path, document_id: str) -> Document:
    """构造与源文件哈希一致的测试 Document。"""

    content = source_path.read_bytes()
    return Document(
        id=document_id,
        user_id=f"user-{document_id}",
        workspace_id=None,
        original_filename=source_path.name,
        content_type=(
            "application/vnd.ms-excel"
            if source_path.suffix.lower() == ".xls"
            else "application/msword"
        ),
        size_bytes=len(content),
        sha256=hashlib.sha256(content).hexdigest(),
    )


def _version(*, document: Document, source_path: Path) -> DocumentVersion:
    """构造与原件哈希一致的明确内容版本，保护派生件谱系边界。"""

    return DocumentVersion(
        document_id=document.id,
        version_number=1,
        storage_tier="UPLOAD",
        storage_path=f"originals/{source_path.name}",
        filename=source_path.name,
        content_type=document.content_type,
        size_bytes=document.size_bytes,
        sha256=document.sha256,
        source_type="UPLOAD",
    )


def _fake_converter(calls: list[list[str]]):
    """返回会生成合法 DOCX 的测试命令执行器。"""

    def run(command: list[str], *, timeout_seconds: int):
        """记录参数并在输出目录创建合法 DOCX。"""

        calls.append(command)
        output_dir = Path(command[command.index("--outdir") + 1])
        converted = DocxDocument()
        converted.add_heading("关于开展测试工作的通知", level=0)
        converted.add_paragraph("这是转换后的正文。")
        converted.save(output_dir / "source.docx")
        return subprocess.CompletedProcess(command, 0, stdout=b"converted", stderr=b"")

    return run


def _fake_xls_converter(calls: list[list[str]]):
    """返回会生成合法多 Sheet XLSX 的确定性 LibreOffice fake。"""

    def run(command: list[str], *, timeout_seconds: int):
        """记录参数并生成供 openpyxl 双重校验的工作簿。"""

        calls.append(command)
        assert Path(command[-1]).name == "source.xls"
        assert any(str(item).startswith("-env:UserInstallation=file:") for item in command)
        output_dir = Path(command[command.index("--outdir") + 1])
        workbook = __import__("openpyxl").Workbook()
        workbook.active.title = "汇总"
        workbook.active.append(["姓名", "金额"])
        workbook.active.append(["张三", 100])
        workbook.create_sheet("明细").append(["学号", "姓名"])
        workbook.save(output_dir / "source.xlsx")
        return subprocess.CompletedProcess(command, 0, stdout=b"converted", stderr=b"")

    return run


def test_resolve_libreoffice_executable_prefers_configured_path(tmp_path):
    """显式配置必须高于 PATH 和平台默认目录。"""

    executable = tmp_path / "custom-soffice"
    executable.write_bytes(b"")

    resolved = resolve_libreoffice_executable(
        configured=str(executable),
        platform_name="linux",
        environ={},
        which=lambda _: None,
    )

    assert resolved == executable


def test_resolve_configured_windows_soffice_exe_prefers_console_sibling(tmp_path):
    """显式配置 GUI exe 时仍应使用同目录可等待的 soffice.com。"""

    program_dir = tmp_path / "LibreOffice" / "program"
    program_dir.mkdir(parents=True)
    executable = program_dir / "soffice.exe"
    console = program_dir / "soffice.com"
    executable.write_bytes(b"")
    console.write_bytes(b"")

    resolved = resolve_libreoffice_executable(
        configured=str(executable),
        platform_name="win32",
        environ={},
        which=lambda _name: None,
    )

    assert resolved == console


def test_resolve_libreoffice_executable_finds_windows_soffice_com_first(tmp_path):
    """Windows 默认目录中必须优先使用 soffice.com。"""

    program_files = tmp_path / "Program Files"
    executable = program_files / "LibreOffice" / "program" / "soffice.com"
    executable.parent.mkdir(parents=True)
    executable.write_bytes(b"")

    resolved = resolve_libreoffice_executable(
        platform_name="win32",
        environ={"ProgramFiles": str(program_files)},
        which=lambda _: None,
    )

    assert resolved == executable


def test_libreoffice_profile_uri_supports_windows_drive_path():
    """Windows LibreOffice profile 必须使用合法 file URI。"""

    assert libreoffice_profile_uri(PureWindowsPath("C:/Temp/file-agent-profile")) == (
        "file:///C:/Temp/file-agent-profile"
    )


def test_doc_conversion_creates_and_reuses_persistent_artifact(monkeypatch, tmp_path):
    """同一 Document 第二次读取必须复用持久派生件。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    source_path = tmp_path / "notice.doc"
    source_path.write_bytes(b"legacy-doc-content")
    executable = tmp_path / "soffice"
    executable.write_bytes(b"")
    db = _session()
    document = _document(source_path=source_path, document_id="document-1")
    db.add(document)
    db.flush()
    calls: list[list[str]] = []
    service = LegacyOfficeConversionService(
        db=db,
        storage_root=tmp_path / "storage",
        executable=executable,
        command_runner=_fake_converter(calls),
        converter_version="LibreOffice Test 1.0",
    )

    first = service.get_or_create_docx(document=document, source_path=source_path)
    second = service.get_or_create_docx(document=document, source_path=source_path)

    assert first.reused is False
    assert second.reused is True
    assert first.storage_path == second.storage_path
    assert first.file_path.exists()
    assert len(calls) == 1
    assert db.query(DocumentArtifact).count() == 1


def test_doc_conversion_stages_output_on_target_volume_before_atomic_replace(
    monkeypatch,
    tmp_path,
):
    """DOCX 派生件必须在目标目录内暂存，不能从 Windows 系统盘跨卷 replace。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    source_path = tmp_path / "notice.doc"
    source_path.write_bytes(b"legacy-doc-cross-volume")
    executable = tmp_path / "soffice"
    executable.write_bytes(b"")
    db = _session()
    document = _document(source_path=source_path, document_id="document-cross-volume")
    db.add(document)
    db.flush()
    real_replace = os.replace
    replace_calls: list[tuple[Path, Path]] = []

    def windows_same_volume_replace(source, target):
        """模拟 Windows：只允许同一目标目录内执行原子替换。"""

        source_path_value = Path(source)
        target_path_value = Path(target)
        assert source_path_value.parent == target_path_value.parent
        replace_calls.append((source_path_value, target_path_value))
        real_replace(source_path_value, target_path_value)

    monkeypatch.setattr(
        "app.modules.files.office_conversion.os.replace",
        windows_same_volume_replace,
    )
    service = LegacyOfficeConversionService(
        db=db,
        storage_root=tmp_path / "storage",
        executable=executable,
        command_runner=_fake_converter([]),
        converter_version="LibreOffice Test 1.0",
    )

    artifact = service.get_or_create_docx(
        document=document,
        source_path=source_path,
    )

    assert artifact.file_path.is_file()
    assert len(replace_calls) == 1
    staged_path, final_path = replace_calls[0]
    assert staged_path.name.endswith(".part")
    assert staged_path.parent == final_path.parent


def test_derivative_staging_name_stays_short_for_long_target_name(tmp_path):
    """临时文件名不能重复长目标名，否则合法目标也会超过 Windows 路径限制。"""

    target_name = f"{'a' * 64}.docx"
    # 最终路径保持在传统 Windows MAX_PATH 内；旧实现把 target_name、PID 和
    # 纳秒时间戳再次拼到临时名后才会越界，短随机名则仍可正常原子提交。
    padding_length = max(
        1,
        245 - len(str(tmp_path.resolve())) - len(target_name) - 2,
    )
    target_parent = tmp_path / ("p" * padding_length)
    target_parent.mkdir()
    source_path = tmp_path / "converted.docx"
    source_path.write_bytes(b"converted-office-content")
    target_path = target_parent / target_name

    _replace_with_retry(source_path, target_path)

    assert target_path.read_bytes() == b"converted-office-content"
    assert not source_path.exists()
    assert not list(tmp_path.rglob(".fa-*.part"))


def test_same_content_across_documents_reuses_physical_artifact(monkeypatch, tmp_path):
    """跨用户同内容应复用物理文件，但保留独立 Artifact 权限记录。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    first_source = tmp_path / "first.doc"
    second_source = tmp_path / "second.doc"
    first_source.write_bytes(b"same-legacy-doc-content")
    second_source.write_bytes(b"same-legacy-doc-content")
    executable = tmp_path / "soffice"
    executable.write_bytes(b"")
    db = _session()
    first_document = _document(source_path=first_source, document_id="document-a")
    second_document = _document(source_path=second_source, document_id="document-b")
    db.add_all([first_document, second_document])
    db.flush()
    calls: list[list[str]] = []
    service = LegacyOfficeConversionService(
        db=db,
        storage_root=tmp_path / "storage",
        executable=executable,
        command_runner=_fake_converter(calls),
        converter_version="LibreOffice Test 1.0",
    )

    first = service.get_or_create_docx(document=first_document, source_path=first_source)
    second = service.get_or_create_docx(document=second_document, source_path=second_source)

    assert first.storage_path == second.storage_path
    assert second.reused is True
    assert len(calls) == 1
    artifacts = db.query(DocumentArtifact).order_by(DocumentArtifact.document_id).all()
    assert [item.document_id for item in artifacts] == ["document-a", "document-b"]
    assert len({item.storage_path for item in artifacts}) == 1


def test_xls_conversion_creates_versioned_persistent_artifact_and_reuses_it(
    monkeypatch,
    tmp_path,
):
    """XLS 首次转换必须持久化 XLSX，后续读取复用且原件字节保持不变。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    source_path = tmp_path / "住宿清单.xls"
    source_path.write_bytes(b"legacy-xls-content")
    original_bytes = source_path.read_bytes()
    executable = tmp_path / "soffice"
    executable.write_bytes(b"")
    db = _session()
    document = _document(source_path=source_path, document_id="document-xls")
    version = _version(document=document, source_path=source_path)
    db.add_all([document, version])
    db.flush()
    calls: list[list[str]] = []
    service = LegacyOfficeConversionService(
        db=db,
        storage_root=tmp_path / "storage",
        executable=executable,
        command_runner=_fake_xls_converter(calls),
        converter_version="LibreOffice Test 1.0",
    )

    first = service.get_or_create_xlsx(
        document=document,
        document_version=version,
        source_path=source_path,
    )
    second = service.get_or_create_xlsx(
        document=document,
        document_version=version,
        source_path=source_path,
    )

    assert first.reused is False
    assert second.reused is True
    assert first.storage_path == second.storage_path
    assert first.file_path.suffix == ".xlsx"
    assert first.file_path.is_file()
    assert source_path.read_bytes() == original_bytes
    assert len(calls) == 1
    artifact = db.query(DocumentArtifact).one()
    assert artifact.artifact_type == "CONVERTED_XLSX"
    assert artifact.document_version_id == version.id
    assert artifact.metadata_json["source_format"] == "xls"
    assert artifact.metadata_json["parsed_format"] == "xlsx"


def test_xls_reuses_valid_persistent_artifact_when_libreoffice_is_unavailable(
    monkeypatch,
    tmp_path,
):
    """转换器临时不可用时仍应复用同规则且通过校验的持久化 XLSX。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    source_path = tmp_path / "住宿清单.xls"
    source_path.write_bytes(b"legacy-xls-content")
    executable = tmp_path / "soffice"
    executable.write_bytes(b"")
    db = _session()
    document = _document(source_path=source_path, document_id="document-xls-offline")
    version = _version(document=document, source_path=source_path)
    db.add_all([document, version])
    db.flush()
    online_service = LegacyOfficeConversionService(
        db=db,
        storage_root=tmp_path / "storage",
        executable=executable,
        command_runner=_fake_xls_converter([]),
        converter_version="LibreOffice Test 1.0",
    )
    created = online_service.get_or_create_xlsx(
        document=document,
        document_version=version,
        source_path=source_path,
    )
    monkeypatch.setattr(
        "app.modules.files.office_conversion.resolve_libreoffice_executable",
        lambda **_: None,
    )

    offline_service = LegacyOfficeConversionService(
        db=db,
        storage_root=tmp_path / "storage",
    )
    reused = offline_service.get_or_create_xlsx(
        document=document,
        document_version=version,
        source_path=source_path,
    )

    assert reused.reused is True
    assert reused.artifact_id == created.artifact_id
    assert reused.converter_config_hash == created.converter_config_hash


def test_xls_conversion_rejects_invalid_persistent_xlsx(monkeypatch, tmp_path):
    """LibreOffice 成功退出但 XLSX 结构无效时不得登记持久化派生件。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    source_path = tmp_path / "损坏表格.xls"
    source_path.write_bytes(b"legacy-xls-content")
    executable = tmp_path / "soffice"
    executable.write_bytes(b"")
    db = _session()
    document = _document(source_path=source_path, document_id="document-xls-invalid")
    version = _version(document=document, source_path=source_path)
    db.add_all([document, version])
    db.flush()

    def invalid_output(command: list[str], *, timeout_seconds: int):
        """生成伪装为 XLSX 的无效输出。"""

        output_dir = Path(command[command.index("--outdir") + 1])
        (output_dir / "source.xlsx").write_bytes(b"not-a-zip")
        return subprocess.CompletedProcess(command, 0, stdout=b"converted", stderr=b"")

    service = LegacyOfficeConversionService(
        db=db,
        storage_root=tmp_path / "storage",
        executable=executable,
        command_runner=invalid_output,
        converter_version="LibreOffice Test 1.0",
    )

    with pytest.raises(OfficeConversionError) as exc_info:
        service.get_or_create_xlsx(
            document=document,
            document_version=version,
            source_path=source_path,
        )

    assert exc_info.value.code == "XLSX_CONVERSION_OUTPUT_INVALID"
    assert db.query(DocumentArtifact).count() == 0


def test_doc_conversion_rejects_missing_output(monkeypatch, tmp_path):
    """LibreOffice 未生成 DOCX 时必须返回稳定错误码。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path / "storage"))
    source_path = tmp_path / "broken.doc"
    source_path.write_bytes(b"broken-content")
    executable = tmp_path / "soffice"
    executable.write_bytes(b"")
    db = _session()
    document = _document(source_path=source_path, document_id="document-broken")
    db.add(document)
    db.flush()

    def no_output(command: list[str], *, timeout_seconds: int):
        """模拟转换命令成功退出但没有生成输出。"""

        return subprocess.CompletedProcess(command, 0, stdout=b"", stderr=b"")

    service = LegacyOfficeConversionService(
        db=db,
        storage_root=tmp_path / "storage",
        executable=executable,
        command_runner=no_output,
        converter_version="LibreOffice Test 1.0",
    )

    with pytest.raises(OfficeConversionError) as exc_info:
        service.get_or_create_docx(document=document, source_path=source_path)

    assert exc_info.value.code == "DOCX_OUTPUT_MISSING"
    assert db.query(DocumentArtifact).count() == 0
