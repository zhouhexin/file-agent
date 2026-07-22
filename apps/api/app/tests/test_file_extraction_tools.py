"""原始文件读取与文本解析 Tool 测试。"""

from __future__ import annotations

import subprocess
from io import BytesIO
from pathlib import Path

from fastapi.testclient import TestClient

from app.core import config
from app.db.models import Document, DocumentArtifact, DocumentElement, DocumentExtractionRun, DocumentPage
from app.modules.files.extractors import extract_document_text
from app.modules.files.extraction_repository import FileExtractionRepository
from app.modules.files.docling_parser import try_parse_with_docling
from app.modules.agent.tool_registry import ToolRegistry
from app.tests.helpers import clear_overrides, client_with_database


class FakeOcrService:
    """测试用 OCR 服务，避免依赖真实 PaddleOCR 或外部 LLM。"""

    def __init__(self, text: str = "OCR 识别文本", source: str = "paddleocr_cpu", quality_score: float = 0.9):
        """保存固定 OCR 返回值。"""

        self.text = text
        self.source = source
        self.quality_score = quality_score
        self.calls: list[dict] = []

    def extract_image(self, *, image_path, page_number: int = 1):
        """记录调用并返回固定 OCR 页面。"""

        self.calls.append({"image_path": image_path, "page_number": page_number})
        return {
            "ok": True,
            "text": self.text,
            "source": self.source,
            "provider_name": self.source,
            "quality_score": self.quality_score,
            "confidence": 0.88,
            "blocks": [],
            "warnings": [],
        }


class _FakeValue:
    """模拟 Docling 枚举值。"""

    def __init__(self, value: str):
        self.value = value


class _FakeBBox:
    """模拟 Docling 边界框模型。"""

    def model_dump(self, mode: str = "json") -> dict:
        assert mode == "json"
        return {"l": 10, "t": 20, "r": 100, "b": 40, "coord_origin": "TOPLEFT"}


class _FakeProvenance:
    page_no = 1
    bbox = _FakeBBox()


class _FakeDoclingItem:
    """模拟一个带位置的 Docling 标题元素。"""

    text = "关于做好测试工作的通知"
    label = _FakeValue("title")
    content_layer = _FakeValue("body")
    prov = [_FakeProvenance()]
    parent = None


class _FakeDoclingDocument:
    """模拟最小 DoclingDocument。"""

    def iterate_items(self):
        yield _FakeDoclingItem(), 1


class _FakeDoclingConverter:
    """模拟本地 Docling converter。"""

    def convert(self, file_path):
        assert file_path.name == "notice.pdf"
        return type("Conversion", (), {"document": _FakeDoclingDocument()})()


def _auth_header(client: TestClient, username: str) -> dict[str, str]:
    """注册并登录测试用户，返回 Authorization header。"""

    client.post(
        "/api/auth/register",
        json={"username": username, "password": "password123", "display_name": username},
    )
    login_response = client.post(
        "/api/auth/login",
        json={"username": username, "password": "password123"},
    )
    return {"Authorization": f"Bearer {login_response.json()['access_token']}"}


def _upload_text(client: TestClient, headers: dict[str, str], filename: str = "notes.txt") -> str:
    """上传一个 UTF-8 文本文件并返回 document_id。"""

    response = client.post(
        "/api/files/upload",
        headers=headers,
        files={"file": (filename, "学生姓名：张三\n奖学金：一等奖\n".encode("utf-8"), "text/plain")},
    )
    assert response.status_code == 202
    return response.json()["document_id"]


def _docx_bytes() -> bytes:
    """构造包含中文正文的 docx 测试文件。"""

    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_paragraph("学生姓名：王五")
    document.add_paragraph("奖学金：三等奖")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _upload_docx(client: TestClient, headers: dict[str, str]) -> str:
    """上传一个 docx 文件并返回 document_id。"""

    response = client.post(
        "/api/files/upload",
        headers=headers,
        files={
            "file": (
                "student.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 202
    return response.json()["document_id"]


def _upload_doc(client: TestClient, headers: dict[str, str]) -> str:
    """上传一个由测试转换器处理的旧版 doc 占位文件。"""

    response = client.post(
        "/api/files/upload",
        headers=headers,
        files={"file": ("legacy-notice.doc", b"legacy-doc-content", "application/msword")},
    )
    assert response.status_code == 202
    return response.json()["document_id"]


def test_extraction_tables_can_be_created():
    """文件解析运行表和页面表必须纳入 ORM metadata。"""

    assert DocumentExtractionRun.__tablename__ == "document_extraction_runs"
    assert DocumentPage.__tablename__ == "document_pages"


def test_read_original_file_returns_metadata_for_owner(monkeypatch, tmp_path):
    """read-original-file 只能读取当前用户自己的文件元信息。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path))
    config.get_settings.cache_clear()
    client, SessionLocal = client_with_database()
    headers = _auth_header(client, "file-reader")
    document_id = _upload_text(client, headers)

    db = SessionLocal()
    try:
        user_id = client.get("/api/auth/me", headers=headers).json()["id"]
        result = ToolRegistry(db=db, user_id=user_id).invoke(
            "read-original-file",
            {"document_id": document_id},
        )

        assert result.status == "COMPLETED"
        assert result.output_json["ok"] is True
        assert result.output_json["document_id"] == document_id
        assert result.output_json["filename"] == "notes.txt"
        assert result.output_json["storage_backend"] == "local"
        assert "storage_path" not in result.output_json
    finally:
        db.close()
        clear_overrides()
        config.get_settings.cache_clear()


def test_read_original_file_rejects_other_users_document(monkeypatch, tmp_path):
    """read-original-file 不能跨用户读取文件。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path))
    config.get_settings.cache_clear()
    client, SessionLocal = client_with_database()
    owner_headers = _auth_header(client, "file-owner")
    other_headers = _auth_header(client, "file-other")
    document_id = _upload_text(client, owner_headers)

    db = SessionLocal()
    try:
        other_user_id = client.get("/api/auth/me", headers=other_headers).json()["id"]
        result = ToolRegistry(db=db, user_id=other_user_id).invoke(
            "read-original-file",
            {"document_id": document_id},
        )

        assert result.output_json["ok"] is False
        assert result.output_json["error"]["code"] == "DOCUMENT_NOT_FOUND"
    finally:
        db.close()
        clear_overrides()
        config.get_settings.cache_clear()


def test_extract_document_text_persists_text_pages(monkeypatch, tmp_path):
    """extract-document-text 应解析文本文件并持久化 DocumentPage。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path))
    config.get_settings.cache_clear()
    client, SessionLocal = client_with_database()
    headers = _auth_header(client, "text-extractor")
    document_id = _upload_text(client, headers)

    db = SessionLocal()
    try:
        user_id = client.get("/api/auth/me", headers=headers).json()["id"]
        result = ToolRegistry(db=db, user_id=user_id).invoke(
            "extract-document-text",
            {"document_id": document_id},
        )

        assert result.output_json["ok"] is True
        assert result.output_json["status"] == "COMPLETED"
        assert result.output_json["read_quality"] == "GOOD"
        assert result.output_json["read_profile"]["file_type"] == "text"
        assert result.output_json["read_profile"]["char_count"] > 0
        assert result.output_json["pages"][0]["char_count"] > 0
        assert db.query(DocumentExtractionRun).count() == 1
        page = db.query(DocumentPage).one()
        assert page.document_id == document_id
        assert "张三" in page.text_content
        assert page.metadata_json["read_quality"] == "GOOD"
    finally:
        db.close()
        clear_overrides()
        config.get_settings.cache_clear()


def test_extract_document_text_persists_docx_pages(monkeypatch, tmp_path):
    """extract-document-text 应解析 docx 正文并持久化 DocumentPage。"""

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path))
    monkeypatch.setenv("DOCLING_ENABLED", "false")
    config.get_settings.cache_clear()
    client, SessionLocal = client_with_database()
    headers = _auth_header(client, "docx-extractor")
    document_id = _upload_docx(client, headers)

    db = SessionLocal()
    try:
        user_id = client.get("/api/auth/me", headers=headers).json()["id"]
        result = ToolRegistry(db=db, user_id=user_id).invoke(
            "extract-document-text",
            {"document_id": document_id},
        )

        assert result.output_json["ok"] is True
        assert result.output_json["status"] == "COMPLETED"
        assert result.output_json["extractor"] == "docx"
        page = db.query(DocumentPage).one()
        assert page.document_id == document_id
        assert "王五" in page.text_content
        assert "三等奖" in page.text_content
    finally:
        db.close()
        clear_overrides()
        config.get_settings.cache_clear()


def test_extract_doc_creates_docx_artifact_and_reuses_it_on_reprocess(monkeypatch, tmp_path):
    """重新解析 DOC 时应复用 DOCX 派生件，而不是重复启动 LibreOffice。"""

    from docx import Document as DocxDocument

    monkeypatch.setenv("FILE_STORAGE_ROOT", str(tmp_path))
    monkeypatch.setenv("DOCLING_ENABLED", "false")
    config.get_settings.cache_clear()
    executable = tmp_path / "soffice"
    executable.write_bytes(b"")
    conversion_calls: list[list[str]] = []

    monkeypatch.setattr(
        "app.modules.files.office_conversion.resolve_libreoffice_executable",
        lambda **_: executable,
    )
    monkeypatch.setattr(
        "app.modules.files.office_conversion.libreoffice_runtime_version",
        lambda _: "LibreOffice Test 1.0",
    )

    def fake_convert(command: list[str], *, timeout_seconds: int):
        """模拟 LibreOffice 在隔离输出目录生成 DOCX。"""

        conversion_calls.append(command)
        output_dir = Path(command[command.index("--outdir") + 1])
        converted = DocxDocument()
        converted.add_heading("关于开展测试工作的通知", level=0)
        converted.add_paragraph("转换后的正文内容。")
        converted.save(output_dir / "source.docx")
        return subprocess.CompletedProcess(command, 0, stdout=b"converted", stderr=b"")

    monkeypatch.setattr(
        "app.modules.files.office_conversion.run_libreoffice_command",
        fake_convert,
    )
    client, SessionLocal = client_with_database()
    headers = _auth_header(client, "legacy-doc-extractor")
    document_id = _upload_doc(client, headers)
    db = SessionLocal()
    try:
        user_id = client.get("/api/auth/me", headers=headers).json()["id"]
        registry = ToolRegistry(db=db, user_id=user_id)

        first = registry.invoke("extract-document-text", {"document_id": document_id})
        second = registry.invoke(
            "extract-document-text",
            {"document_id": document_id, "force_reprocess": True},
        )

        assert first.output_json["ok"] is True
        assert first.output_json["conversion_reused"] is False
        assert first.output_json["conversion_source_format"] == "doc"
        assert first.output_json["conversion_parsed_format"] == "docx"
        assert first.output_json["conversion_converter"] == "libreoffice"
        assert first.output_json["conversion_converter_version"] == "LibreOffice Test 1.0"
        assert second.output_json["ok"] is True
        assert second.output_json["conversion_reused"] is True
        assert len(conversion_calls) == 1
        assert db.query(DocumentArtifact).count() == 1
        assert db.query(DocumentExtractionRun).count() == 2
        assert db.query(DocumentPage).count() == 2
        latest_page = db.query(DocumentPage).order_by(DocumentPage.created_at.desc()).first()
        assert latest_page.metadata_json["source_format"] == "doc"
        assert latest_page.metadata_json["parsed_format"] == "docx"
        assert latest_page.metadata_json["conversion_reused"] is True
        assert "关于开展测试工作的通知" in latest_page.text_content
    finally:
        db.close()
        clear_overrides()
        config.get_settings.cache_clear()


def test_extract_document_text_supports_doc_with_textutil(monkeypatch, tmp_path):
    """extract-document-text 应通过系统转换器读取旧版 doc 正文。"""

    doc_path = tmp_path / "promise.doc"
    doc_path.write_bytes(b"legacy-doc")

    def fake_which(name: str) -> str | None:
        """测试中只模拟 textutil 可用，避免依赖本机 Office 工具。"""

        return "/usr/bin/textutil" if name == "textutil" else None

    def fake_run(*args, **kwargs) -> subprocess.CompletedProcess[bytes]:
        """模拟 textutil 将 doc 转成纯文本输出。"""

        return subprocess.CompletedProcess(args=args[0], returncode=0, stdout="电子发票承诺书\n经办人：李四".encode("utf-8"), stderr=b"")

    monkeypatch.setattr("app.modules.files.extractors.shutil.which", fake_which)
    monkeypatch.setattr("app.modules.files.extractors.subprocess.run", fake_run)

    result = extract_document_text(
        file_path=doc_path,
        filename="电子发票承诺书.doc",
        content_type="application/msword",
    )

    assert result["ok"] is True
    assert result["status"] == "COMPLETED"
    assert result["extractor"] == "doc-textutil"
    assert "电子发票承诺书" in result["pages"][0]["text"]


def test_extract_document_text_prefers_docling_for_pdf(monkeypatch, tmp_path):
    """默认开启 Docling 时，PDF 应优先返回结构化页面和元素。"""

    monkeypatch.setenv("DOCLING_ENABLED", "true")
    config.get_settings.cache_clear()
    pdf_path = tmp_path / "notice.pdf"
    pdf_path.write_bytes(b"fake-pdf")
    monkeypatch.setattr(
        "app.modules.files.extractors.try_parse_with_docling",
        lambda **kwargs: {
            "ok": True,
            "extractor": "docling",
            "pages": [
                {
                    "page_number": 1,
                    "sheet_name": None,
                    "text": "关于做好测试工作的通知",
                    "metadata": {"structured": True},
                }
            ],
            "elements": [
                {
                    "element_index": 0,
                    "label": "title",
                    "text": "关于做好测试工作的通知",
                    "page_number": 1,
                    "bbox": {"l": 10, "t": 20, "r": 100, "b": 40},
                    "content_layer": "body",
                    "parent_ref": None,
                    "metadata": {},
                }
            ],
            "warnings": [],
        },
    )

    result = extract_document_text(
        file_path=pdf_path,
        filename="notice.pdf",
        content_type="application/pdf",
    )

    assert result["ok"] is True
    assert result["extractor"] == "docling"
    assert result["elements"][0]["label"] == "title"
    assert result["pages"][0]["metadata"]["structured"] is True


def test_extract_document_text_falls_back_when_docling_fails(monkeypatch, tmp_path):
    """Docling 不可用或转换失败时必须继续使用现有 PDF 解析器。"""

    monkeypatch.setenv("DOCLING_ENABLED", "true")
    config.get_settings.cache_clear()
    pdf_path = tmp_path / "fallback.pdf"
    pdf_path.write_bytes(b"fake-pdf")
    monkeypatch.setattr(
        "app.modules.files.extractors.try_parse_with_docling",
        lambda **kwargs: {
            "ok": False,
            "error": {"code": "DOCLING_NOT_AVAILABLE", "message": "未安装 Docling"},
        },
    )
    monkeypatch.setattr(
        "app.modules.files.extractors._extract_pdf_native_pages",
        lambda file_path: [
            {
                "page_number": 1,
                "sheet_name": None,
                "text": "现有解析器正文",
                "metadata": {"page_index": 0},
            }
        ],
    )

    result = extract_document_text(
        file_path=pdf_path,
        filename="fallback.pdf",
        content_type="application/pdf",
    )

    assert result["ok"] is True
    assert result["extractor"] == "pdf"
    assert result["pages"][0]["text"] == "现有解析器正文"
    assert result["warnings"][0]["code"] == "DOCLING_NOT_AVAILABLE"


def test_docling_adapter_serializes_document_elements(monkeypatch, tmp_path):
    """Docling 适配器应输出项目稳定的页面、标签和位置结构。"""

    pdf_path = tmp_path / "notice.pdf"
    pdf_path.write_bytes(b"fake-pdf")
    monkeypatch.setattr(
        "app.modules.files.docling_parser._build_converter",
        lambda ocr_enabled: _FakeDoclingConverter(),
    )

    result = try_parse_with_docling(
        file_path=pdf_path,
        filename=pdf_path.name,
        content_type="application/pdf",
    )

    assert result["ok"] is True
    assert result["elements"][0]["label"] == "title"
    assert result["elements"][0]["page_number"] == 1
    assert result["elements"][0]["bbox"]["t"] == 20
    assert result["pages"][0]["text"] == "关于做好测试工作的通知"


def test_extraction_repository_persists_structured_elements(tmp_path, monkeypatch):
    """结构化元素必须和解析运行一起写入并可按配置指纹复用。"""

    monkeypatch.chdir(tmp_path)
    client, SessionLocal = client_with_database()
    headers = _auth_header(client, "structured-owner")
    document_id = _upload_text(client, headers, filename="structured.txt")
    db = SessionLocal()
    try:
        document = db.get(Document, document_id)
        assert document is not None
        repository = FileExtractionRepository(db, document.user_id)
        run = repository.create_extraction_run(
            document_id=document.id,
            extractor="docling@2.50.0",
            parser_name="docling",
            parser_version="2.50.0",
            parser_config_hash="a" * 64,
        )
        repository.complete_extraction_run(
            run=run,
            pages=[{"page_number": 1, "text": "结构化正文", "metadata": {}}],
            elements=[
                {
                    "element_index": 0,
                    "label": "title",
                    "text": "结构化标题",
                    "page_number": 1,
                    "bbox": {"l": 1, "t": 2, "r": 3, "b": 4},
                    "content_layer": "body",
                    "metadata": {"hierarchy_level": 1},
                }
            ],
        )

        reusable = repository.get_latest_successful_extraction(
            document_id=document.id,
            parser_config_hash="a" * 64,
        )

        assert reusable is not None
        assert reusable["run"].parser_name == "docling"
        assert reusable["elements"][0].label == "title"
        assert reusable["elements"][0].bbox_json["t"] == 2
        assert db.query(DocumentElement).count() == 1
    finally:
        db.close()
        clear_overrides()


def test_extract_document_text_always_converts_xls_before_reading(monkeypatch, tmp_path):
    """旧版 XLS 必须先转换临时 XLSX，并完整读取所有工作表。"""

    xls_path = tmp_path / "奖学金汇总.xls"
    xls_path.write_bytes(b"legacy-xls")
    original_bytes = xls_path.read_bytes()

    def fake_converter(*, source_path, output_dir):
        """模拟 LibreOffice 在隔离目录生成包含多工作表的临时文件。"""

        assert source_path == xls_path
        converted_path = output_dir / "source.xlsx"
        workbook = __import__("openpyxl").Workbook()
        worksheet = workbook.active
        worksheet.title = "汇总"
        worksheet.append(["姓名", "等级"])
        worksheet.append(["赵六", "一等奖"])
        detail_sheet = workbook.create_sheet("明细")
        detail_sheet.append(["学号", "姓名"])
        detail_sheet.append(["2026001", "赵六"])
        workbook.save(converted_path)
        return converted_path

    monkeypatch.setattr("app.modules.files.extractors.convert_xls_to_xlsx", fake_converter)

    result = extract_document_text(
        file_path=xls_path,
        filename="奖学金汇总.xls",
        content_type="application/vnd.ms-excel",
    )

    assert result["ok"] is True
    assert result["status"] == "COMPLETED"
    assert result["extractor"] == "excel-xls-converted"
    assert [page["sheet_name"] for page in result["pages"]] == ["汇总", "明细"]
    # 临时 XLSX 的每个非空行必须携带真实坐标，供阶段三 Evidence 精确引用。
    assert result["pages"][0]["metadata"]["line_cell_ranges"] == [
        {"line_index": 0, "row_number": 1, "cell_range": "A1:B1"},
        {"line_index": 1, "row_number": 2, "cell_range": "A2:B2"},
    ]
    assert result["pages"][1]["metadata"]["used_cell_range"] == "A1:B2"
    assert "赵六\t一等奖" in result["pages"][0]["text"]
    assert result["pages"][0]["metadata"]["converted_from"] == ".xls"
    assert xls_path.read_bytes() == original_bytes


def test_extract_document_text_returns_structured_xls_conversion_failure(monkeypatch, tmp_path):
    """LibreOffice 转换失败必须保留稳定错误码，不能回退为文件名正文。"""

    xls_path = tmp_path / "损坏表格.xls"
    xls_path.write_bytes(b"broken-xls")

    def fail_conversion(**_kwargs):
        """模拟转换器不可用。"""

        from app.modules.spreadsheet_analysis.conversion import SpreadsheetConversionError

        raise SpreadsheetConversionError("XLS_CONVERTER_NOT_AVAILABLE", "未找到 LibreOffice。")

    monkeypatch.setattr("app.modules.files.extractors.convert_xls_to_xlsx", fail_conversion)

    result = extract_document_text(
        file_path=xls_path,
        filename=xls_path.name,
        content_type="application/vnd.ms-excel",
    )

    assert result["ok"] is False
    assert result["status"] == "FAILED"
    assert result["error"]["code"] == "XLS_CONVERTER_NOT_AVAILABLE"


def test_extract_image_uses_injected_ocr_service(monkeypatch, tmp_path):
    """图片解析应通过 OCR 服务写入统一页面文本。"""

    # 测试套件默认关闭真实 OCR；本用例只显式启用注入的 deterministic fake。
    monkeypatch.setenv("OCR_ENABLED", "true")
    config.get_settings.cache_clear()
    image_path = tmp_path / "scan.png"
    image_path.write_bytes(b"fake-image")
    ocr_service = FakeOcrService(text="电子发票承诺书 OCR 文本")

    result = extract_document_text(
        file_path=image_path,
        filename="scan.png",
        content_type="image/png",
        ocr_service=ocr_service,
    )

    assert result["ok"] is True
    assert result["extractor"] == "paddleocr_cpu"
    assert result["pages"][0]["text"] == "电子发票承诺书 OCR 文本"
    assert result["pages"][0]["metadata"]["ocr_source"] == "paddleocr_cpu"
    assert ocr_service.calls[0]["page_number"] == 1


def test_empty_pdf_triggers_ocr_fallback(monkeypatch, tmp_path):
    """PDF 原生文本为空时应渲染页面并触发 OCR 兜底。"""

    # 只允许注入的 fake OCR 执行，不能在 Windows 测试机下载或初始化 Paddle 模型。
    monkeypatch.setenv("OCR_ENABLED", "true")
    config.get_settings.cache_clear()
    pdf_path = tmp_path / "scan.pdf"
    pdf_path.write_bytes(b"fake-pdf")
    rendered_page = tmp_path / "page-1.png"
    rendered_page.write_bytes(b"fake-render")
    ocr_service = FakeOcrService(text="扫描 PDF OCR 文本")

    monkeypatch.setattr(
        "app.modules.files.extractors._extract_pdf_native_pages",
        lambda file_path: [{"page_number": 1, "sheet_name": None, "text": "", "metadata": {"page_index": 0}}],
    )
    monkeypatch.setattr(
        "app.modules.files.extractors._render_pdf_pages_for_ocr",
        lambda file_path, page_numbers: {1: rendered_page},
    )

    result = extract_document_text(
        file_path=pdf_path,
        filename="scan.pdf",
        content_type="application/pdf",
        ocr_service=ocr_service,
    )

    assert result["ok"] is True
    assert result["extractor"] == "pdf+paddleocr_cpu"
    assert result["pages"][0]["text"] == "扫描 PDF OCR 文本"
    assert result["pages"][0]["metadata"]["ocr_fallback"] is True
    assert ocr_service.calls[0]["image_path"] == rendered_page


def test_empty_pdf_marks_ocr_needed_when_ocr_disabled(monkeypatch, tmp_path):
    """PDF 原生文本为空且 OCR 关闭时，应返回统一读取质量而不是误报 GOOD。"""

    monkeypatch.setenv("OCR_ENABLED", "false")
    config.get_settings.cache_clear()
    pdf_path = tmp_path / "scan.pdf"
    pdf_path.write_bytes(b"fake-pdf")
    monkeypatch.setattr(
        "app.modules.files.extractors._extract_pdf_native_pages",
        lambda file_path: [{"page_number": 1, "sheet_name": None, "text": "", "metadata": {"page_index": 0}}],
    )

    result = extract_document_text(
        file_path=pdf_path,
        filename="scan.pdf",
        content_type="application/pdf",
    )

    assert result["ok"] is True
    assert result["read_quality"] == "OCR_NEEDED"
    assert result["read_profile"]["requires_ocr"] is True
    assert result["read_profile"]["char_count"] == 0
    assert result["pages"][0]["metadata"]["read_quality"] == "OCR_NEEDED"
    config.get_settings.cache_clear()
