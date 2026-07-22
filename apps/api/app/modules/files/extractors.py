"""上传文件内容解析器。"""

from __future__ import annotations

import csv
import hashlib
import shutil
import subprocess
import tempfile
from io import StringIO
from pathlib import Path
from typing import Any, Dict, List

from app.core.config import get_settings
from app.modules.files.docling_parser import docling_runtime_version, try_parse_with_docling
from app.modules.ocr.service import build_default_ocr_service
from app.modules.spreadsheet_analysis.conversion import SpreadsheetConversionError, convert_xls_to_xlsx


def extract_document_text(*, file_path: Path, filename: str, content_type: str, ocr_service: Any = None) -> Dict[str, Any]:
    """按文件类型解析文本内容，并返回统一结构。"""

    suffix = Path(filename).suffix.lower()
    parser_config_hash = extraction_config_hash(filename=filename)
    docling_failure = _try_docling_first(
        file_path=file_path,
        filename=filename,
        content_type=content_type,
        suffix=suffix,
    )
    if docling_failure.get("ok"):
        return _completed(
            docling_failure["extractor"],
            docling_failure["pages"],
            elements=docling_failure.get("elements", []),
            warnings=docling_failure.get("warnings", []),
            parser_name="docling",
            parser_version=docling_runtime_version(),
            parser_config_hash=parser_config_hash or "",
        )
    native_result = extract_document_text_native(
        file_path=file_path,
        filename=filename,
        content_type=content_type,
        ocr_service=ocr_service,
    )
    if suffix in {".docx", ".pdf"}:
        return _apply_parser_metadata(
            _append_parser_warning(native_result, docling_failure),
            parser_config_hash=parser_config_hash,
        )
    return native_result


def extract_document_text_native(
    *,
    file_path: Path,
    filename: str,
    content_type: str,
    ocr_service: Any = None,
) -> Dict[str, Any]:
    """绕过 Docling 并按文件类型调用项目原生解析器。"""

    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"} or content_type.startswith("text/"):
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        return _completed("plain-text", [{"page_number": 1, "sheet_name": None, "text": text, "metadata": {}}])
    if suffix == ".csv":
        text = _extract_csv_text(file_path)
        return _completed("csv", [{"page_number": 1, "sheet_name": None, "text": text, "metadata": {}}])
    if suffix == ".xls":
        return _extract_legacy_xls_text(file_path)
    if suffix == ".xlsx":
        return _extract_excel_text(file_path)
    if suffix == ".doc" or content_type == "application/msword":
        return _extract_doc_text(file_path)
    if suffix == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx_text(file_path)
    if suffix == ".pdf" or content_type == "application/pdf":
        return _extract_pdf_text(file_path, ocr_service=ocr_service)
    if content_type.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}:
        return _extract_image_text(file_path, ocr_service=ocr_service)
    return _failed("unsupported", "UNSUPPORTED_FILE_TYPE", f"暂不支持解析该文件类型：{filename}")


def _try_docling_first(*, file_path: Path, filename: str, content_type: str, suffix: str) -> Dict[str, Any]:
    """对配置格式优先调用 Docling，其余格式返回未启用占位。"""

    settings = get_settings()
    normalized_suffix = suffix.lstrip(".")
    if not settings.docling_enabled or normalized_suffix not in set(settings.docling_formats):
        return {"ok": False, "skipped": True}
    return try_parse_with_docling(
        file_path=file_path,
        filename=filename,
        content_type=content_type,
        ocr_enabled=settings.docling_ocr_enabled,
    )


def extraction_config_hash(*, filename: str) -> str | None:
    """返回影响结构化解析复用的稳定配置指纹。"""

    settings = get_settings()
    suffix = Path(filename).suffix.lower().lstrip(".")
    if not settings.docling_enabled or suffix not in set(settings.docling_formats):
        return None
    identity = "|".join(
        [
            "docling-preferred",
            docling_runtime_version(),
            f"ocr={int(settings.docling_ocr_enabled)}",
            f"format={suffix}",
        ]
    )
    return hashlib.sha256(identity.encode("utf-8")).hexdigest()


def _append_parser_warning(result: Dict[str, Any], docling_result: Dict[str, Any]) -> Dict[str, Any]:
    """把 Docling 回退原因附加到现有解析结果，不改变原结果状态。"""

    error = docling_result.get("error") or {}
    if not error:
        return result
    result["warnings"] = [
        *result.get("warnings", []),
        {
            "code": str(error.get("code") or "DOCLING_FALLBACK"),
            "message": str(error.get("message") or "Docling 未生成可用结果，已使用现有解析器。"),
        },
    ]
    return result


def _apply_parser_metadata(result: Dict[str, Any], *, parser_config_hash: str | None) -> Dict[str, Any]:
    """标记结构化优先策略下实际使用的回退解析器。"""

    if parser_config_hash is None:
        return result
    result["parser_name"] = str(result.get("extractor") or "legacy")
    result["parser_version"] = ""
    result["parser_config_hash"] = parser_config_hash
    return result


def _extract_csv_text(file_path: Path) -> str:
    """使用标准库读取 CSV 并转成行文本。"""

    raw_text = file_path.read_text(encoding="utf-8", errors="ignore")
    rows = csv.reader(StringIO(raw_text))
    return "\n".join(["\t".join(row) for row in rows])


def _extract_excel_text(file_path: Path) -> Dict[str, Any]:
    """使用 openpyxl 读取 Excel 文本。"""

    try:
        import openpyxl
    except ImportError:
        return _failed("excel", "EXCEL_EXTRACTOR_NOT_AVAILABLE", "缺少 openpyxl，无法解析 Excel 文件。")

    workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    pages: List[Dict[str, Any]] = []
    for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
        lines = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                lines.append("\t".join(values))
        pages.append(
            {
                "page_number": sheet_index,
                "sheet_name": sheet.title,
                "text": "\n".join(lines),
                "metadata": {"sheet_index": sheet_index},
            }
        )
    workbook.close()
    return _completed("excel", pages)


def _extract_legacy_xls_text(file_path: Path) -> Dict[str, Any]:
    """把旧版 XLS 隔离转换为临时 XLSX 后再读取，禁止直接解析原件。"""

    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            converted_path = convert_xls_to_xlsx(
                source_path=file_path,
                output_dir=Path(temp_dir),
            )
        except SpreadsheetConversionError as exc:
            return _failed(
                "excel-xls",
                exc.code,
                exc.message,
            )

        result = _extract_excel_text(converted_path)

    if not result.get("ok"):
        result["extractor"] = "excel-xls-converted"
        return result

    result["extractor"] = "excel-xls-converted"
    for page in result.get("pages", []):
        metadata = page.setdefault("metadata", {})
        metadata["converted_from"] = ".xls"
        metadata["converter"] = "libreoffice"
    return result


def _extract_docx_text(file_path: Path) -> Dict[str, Any]:
    """使用 python-docx 读取 docx 段落和表格文本。"""

    try:
        from docx import Document as DocxDocument
    except ImportError:
        return _failed("docx", "DOCX_EXTRACTOR_NOT_AVAILABLE", "缺少 python-docx，无法解析 docx 文件。")

    document = DocxDocument(file_path)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    lines = [paragraph.text for paragraph in paragraphs]
    elements = []
    for element_index, paragraph in enumerate(paragraphs):
        style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")
        normalized_style = style_name.lower()
        label = "title" if "title" in normalized_style else "section_header" if "heading" in normalized_style else "paragraph"
        font_sizes = [run.font.size.pt for run in paragraph.runs if run.font.size is not None]
        bold_runs = [run for run in paragraph.runs if run.text.strip()]
        elements.append(
            {
                "element_index": element_index,
                "label": label,
                "text": paragraph.text,
                "page_number": 1,
                "bbox": None,
                "content_layer": "body",
                "parent_ref": None,
                "metadata": {
                    "style_name": style_name,
                    "alignment": str(paragraph.alignment or ""),
                    "max_font_size": max(font_sizes) if font_sizes else None,
                    "bold_ratio": (
                        sum(1 for run in bold_runs if run.bold is True) / len(bold_runs)
                        if bold_runs
                        else 0
                    ),
                },
            }
        )
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                lines.append("\t".join(values))
    return _completed(
        "docx",
        [
            {
                "page_number": 1,
                "sheet_name": None,
                "text": "\n".join(lines),
                "metadata": {"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)},
            }
        ],
        elements=elements,
    )


def _extract_doc_text(file_path: Path) -> Dict[str, Any]:
    """读取旧版 Word doc 文件，优先使用系统转换工具抽取正文。"""

    textutil_result = _extract_doc_text_with_textutil(file_path)
    if textutil_result is not None:
        return textutil_result

    libreoffice_result = _extract_doc_text_with_libreoffice(file_path)
    if libreoffice_result is not None:
        return libreoffice_result

    return _failed(
        "doc",
        "DOC_CONVERTER_NOT_AVAILABLE",
        "缺少可用的 doc 转换工具，无法解析旧版 Word 文件。请在服务器安装 LibreOffice，或将文件另存为 docx 后重新上传。",
    )


def _extract_doc_text_with_textutil(file_path: Path) -> Dict[str, Any] | None:
    """在 macOS 环境下通过 textutil 将 doc 转成纯文本。"""

    if not shutil.which("textutil"):
        return None

    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(file_path)],
            capture_output=True,
            check=False,
            timeout=30,
        )
    except (OSError, subprocess.SubprocessError) as exc:
        return _failed("doc-textutil", "DOC_TEXTUTIL_FAILED", f"textutil 解析 doc 失败：{exc}")

    if result.returncode != 0:
        error_message = result.stderr.decode("utf-8", errors="ignore").strip()
        return _failed("doc-textutil", "DOC_TEXTUTIL_FAILED", f"textutil 解析 doc 失败：{error_message or '未知错误'}")

    text = result.stdout.decode("utf-8", errors="ignore")
    return _completed(
        "doc-textutil",
        [{"page_number": 1, "sheet_name": None, "text": text, "metadata": {"converter": "textutil"}}],
    )


def _extract_doc_text_with_libreoffice(file_path: Path) -> Dict[str, Any] | None:
    """在服务器环境下通过 LibreOffice 将 doc 转成 txt 后读取。"""

    converter = shutil.which("soffice") or shutil.which("libreoffice")
    if not converter:
        return None

    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            result = subprocess.run(
                [
                    converter,
                    "--headless",
                    "--convert-to",
                    "txt",
                    "--outdir",
                    temp_dir,
                    str(file_path),
                ],
                capture_output=True,
                check=False,
                timeout=60,
            )
        except (OSError, subprocess.SubprocessError) as exc:
            return _failed("doc-libreoffice", "DOC_LIBREOFFICE_FAILED", f"LibreOffice 解析 doc 失败：{exc}")

        if result.returncode != 0:
            error_message = result.stderr.decode("utf-8", errors="ignore").strip()
            return _failed("doc-libreoffice", "DOC_LIBREOFFICE_FAILED", f"LibreOffice 解析 doc 失败：{error_message or '未知错误'}")

        text_files = sorted(Path(temp_dir).glob("*.txt"))
        if not text_files:
            return _failed("doc-libreoffice", "DOC_LIBREOFFICE_OUTPUT_MISSING", "LibreOffice 未生成 doc 文本结果。")

        text = text_files[0].read_text(encoding="utf-8", errors="ignore")
        return _completed(
            "doc-libreoffice",
            [{"page_number": 1, "sheet_name": None, "text": text, "metadata": {"converter": "libreoffice"}}],
        )


def _extract_pdf_text(file_path: Path, ocr_service: Any = None) -> Dict[str, Any]:
    """优先读取 PDF 原生文本；全文为空时按页渲染并进入 OCR。"""

    try:
        native_pages = _extract_pdf_native_pages(file_path)
    except RuntimeError as exc:
        return _failed("pdf", "PDF_EXTRACTOR_NOT_AVAILABLE", str(exc))
    if any(page.get("text", "").strip() for page in native_pages):
        return _completed("pdf", native_pages)

    if not _ocr_enabled():
        return _completed("pdf", native_pages)

    service = ocr_service or build_default_ocr_service()
    try:
        rendered_pages = _render_pdf_pages_for_ocr(
            file_path=file_path,
            page_numbers=[int(page["page_number"]) for page in native_pages],
        )
    except RuntimeError as exc:
        return _failed("pdf-ocr", "PDF_RENDER_FOR_OCR_FAILED", str(exc))
    ocr_pages: List[Dict[str, Any]] = []
    extractor_name = "pdf+ocr"
    for page in native_pages:
        page_number = int(page["page_number"])
        rendered_path = rendered_pages.get(page_number)
        if rendered_path is None:
            ocr_pages.append(page)
            continue
        ocr_result = service.extract_image(image_path=rendered_path, page_number=page_number)
        if not ocr_result.get("ok"):
            ocr_pages.append(
                {
                    **page,
                    "metadata": {
                        **page.get("metadata", {}),
                        "ocr_fallback": True,
                        "ocr_error": ocr_result.get("error"),
                    },
                }
            )
            continue
        extractor_name = f"pdf+{ocr_result.get('source') or 'ocr'}"
        ocr_pages.append(_page_from_ocr_result(page_number=page_number, ocr_result=ocr_result, base_metadata=page.get("metadata", {})))
    return _completed(extractor_name, ocr_pages)


def _extract_pdf_native_pages(file_path: Path) -> List[Dict[str, Any]]:
    """使用 PyMuPDF 读取 PDF 页面原生文本。"""

    try:
        import fitz
    except ImportError:
        raise RuntimeError("缺少 PyMuPDF，无法解析 PDF 文件。") from None

    pages: List[Dict[str, Any]] = []
    with fitz.open(file_path) as document:
        for index, page in enumerate(document, start=1):
            pages.append(
                {
                    "page_number": index,
                    "sheet_name": None,
                    "text": page.get_text("text"),
                    "metadata": {"page_index": index - 1},
                }
            )
    return pages


def _render_pdf_pages_for_ocr(*, file_path: Path, page_numbers: List[int]) -> Dict[int, Path]:
    """把需要 OCR 的 PDF 页面渲染为临时 PNG 图片。"""

    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("缺少 PyMuPDF，无法渲染 PDF 页面进行 OCR。") from exc

    output_dir = Path(tempfile.mkdtemp(prefix="file-agent-pdf-ocr-"))
    rendered_pages: Dict[int, Path] = {}
    wanted = set(page_numbers)
    with fitz.open(file_path) as document:
        for index, page in enumerate(document, start=1):
            if index not in wanted:
                continue
            output_path = output_dir / f"page-{index:04d}.png"
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            pixmap.save(output_path)
            rendered_pages[index] = output_path
    return rendered_pages


def _extract_image_text(file_path: Path, ocr_service: Any = None) -> Dict[str, Any]:
    """使用 OCR 服务对图片做文字识别。"""

    if not _ocr_enabled():
        return _failed("ocr", "OCR_DISABLED", "OCR 未启用，无法解析图片文字。")

    try:
        service = ocr_service or build_default_ocr_service()
        ocr_result = service.extract_image(image_path=file_path, page_number=1)
    except Exception as exc:
        return _failed("ocr", "OCR_ENGINE_NOT_AVAILABLE", f"OCR 引擎不可用：{exc}")
    if not ocr_result.get("ok"):
        error = ocr_result.get("error") or {}
        return _failed("ocr", str(error.get("code") or "OCR_FAILED"), str(error.get("message") or "OCR 识别失败。"))
    return _completed(
        str(ocr_result.get("source") or "ocr"),
        [_page_from_ocr_result(page_number=1, ocr_result=ocr_result, base_metadata={})],
    )


def _page_from_ocr_result(*, page_number: int, ocr_result: Dict[str, Any], base_metadata: Dict[str, Any]) -> Dict[str, Any]:
    """把 OCR 结果转成 document_pages 可持久化的页面结构。"""

    return {
        "page_number": page_number,
        "sheet_name": None,
        "text": str(ocr_result.get("text") or ""),
        "metadata": {
            **base_metadata,
            "ocr_fallback": True,
            "ocr_source": ocr_result.get("source"),
            "ocr_provider": ocr_result.get("provider_name"),
            "ocr_quality_score": ocr_result.get("quality_score"),
            "ocr_confidence": ocr_result.get("confidence"),
            "ocr_is_llm_fallback": bool(ocr_result.get("is_fallback")),
            "ocr_blocks": ocr_result.get("blocks") or [],
            "ocr_warnings": ocr_result.get("warnings") or [],
        },
    }


def _ocr_enabled() -> bool:
    """读取 OCR 开关，便于测试和部署控制。"""

    return get_settings().ocr_enabled


def _completed(
    extractor: str,
    pages: List[Dict[str, Any]],
    *,
    elements: List[Dict[str, Any]] | None = None,
    warnings: List[Dict[str, Any]] | None = None,
    parser_name: str = "",
    parser_version: str = "",
    parser_config_hash: str = "",
) -> Dict[str, Any]:
    """构造解析成功结果。"""

    read_profile = _build_read_profile(extractor=extractor, pages=pages)
    read_quality = _read_quality_from_profile(profile=read_profile)
    profiled_pages = [
        {
            **page,
            "metadata": {
                **page.get("metadata", {}),
                "read_quality": read_quality,
            },
        }
        for page in pages
    ]
    return {
        "ok": True,
        "status": "COMPLETED",
        "extractor": extractor,
        "read_quality": read_quality,
        "read_profile": read_profile,
        "pages": profiled_pages,
        "elements": elements or [],
        "warnings": warnings or [],
        "parser_name": parser_name,
        "parser_version": parser_version,
        "parser_config_hash": parser_config_hash,
    }


def _failed(extractor: str, code: str, message: str) -> Dict[str, Any]:
    """构造结构化解析失败结果。"""

    return {
        "ok": False,
        "status": "FAILED",
        "extractor": extractor,
        "read_quality": "FAILED",
        "read_profile": {
            "file_type": _file_type_from_extractor(extractor),
            "page_count": 0,
            "sheet_count": 0,
            "char_count": 0,
            "has_text": False,
            "requires_ocr": False,
            "ocr_used": False,
        },
        "error": {
            "code": code,
            "message": message,
            "retryable": False,
            "user_action_required": False,
        },
        "pages": [],
    }


def _build_read_profile(*, extractor: str, pages: List[Dict[str, Any]]) -> Dict[str, Any]:
    """基于解析页生成统一读取 Profile。"""

    char_count = sum(len(str(page.get("text") or "")) for page in pages)
    sheet_count = len([page for page in pages if page.get("sheet_name")])
    ocr_used = any(bool((page.get("metadata") or {}).get("ocr_fallback")) for page in pages) or _extractor_uses_ocr(extractor)
    requires_ocr = extractor == "pdf" and char_count == 0 and bool(pages)
    return {
        "file_type": _file_type_from_extractor(extractor),
        "page_count": len(pages),
        "sheet_count": sheet_count,
        "char_count": char_count,
        "has_text": char_count > 0,
        "requires_ocr": requires_ocr,
        "ocr_used": ocr_used,
    }


def _read_quality_from_profile(*, profile: Dict[str, Any]) -> str:
    """把 Profile 转换为前端和 Agent 可直接使用的读取质量枚举。"""

    if profile.get("requires_ocr"):
        return "OCR_NEEDED"
    if not profile.get("has_text"):
        return "PARTIAL"
    return "GOOD"


def _file_type_from_extractor(extractor: str) -> str:
    """把具体解析器名称归一成文件类型。"""

    if extractor in {"plain-text", "csv"}:
        return "text" if extractor == "plain-text" else "spreadsheet"
    if extractor == "excel":
        return "spreadsheet"
    if extractor.startswith("doc"):
        return "document"
    if extractor.startswith("pdf"):
        return "pdf"
    if extractor in {"ocr", "paddleocr_cpu", "llm_ocr_remote"}:
        return "image"
    return "unknown"


def _extractor_uses_ocr(extractor: str) -> bool:
    """判断解析器是否已经使用 OCR。"""

    return "ocr" in extractor or extractor in {"paddleocr_cpu", "llm_ocr_remote"}
